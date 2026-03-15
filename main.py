import os
from docx import Document
from pdf_parser import extract_text_from_pdf
from translator_engine import AITranslator


def create_bilingual_word(input_pdf, output_docx):
    """
        核心业务流水线：提取 PDF -> 分段 AI 翻译 -> 写入 Word 文档
    """
    print("=" * 40)
    print("🚀 启动离线 AI 翻译流水线")
    print("=" * 40)

    # 1.提取PDF文本
    pages_text = extract_text_from_pdf(input_pdf)
    if not pages_text:
        print("[!] 未提取到文本，流水线终止。")
        return

    # 2.初始化AI引擎
    engine = AITranslator()

    # 3.初始化空白的word文档
    doc = Document()
    doc.add_heading('AI 离线双语翻译报告', 0)

    # 4.逐页逐段翻译
    for page_num, pages_content in enumerate(pages_text):
        print(f"[*] 正在翻译第 {page_num + 1}/{len(pages_text)} 页...")
        doc.add_heading(f'--- 第 {page_num + 1} 页 ---', level=1)

        # 按照换行符切分段落，避免超过AI模型最大Token
        paragraphs = pages_content.split('\n')

        for para in paragraphs:
            if isinstance(para, list):
                para = " ".join([str(item) for item in para])

            para = str(para).strip()

            if not para:
                continue

            # 调用AI引擎进行翻译
            translated_para = str(engine.translate(para))

            # 写入word文档
            p = doc.add_paragraph()
            p.add_run(para + '\n').bold = True
            p.add_run(translated_para)

    os.makedirs(os.path.dirname(output_docx), exist_ok=True)
    doc.save(output_docx)

    print("\n" + "=" * 40)
    print(f"[+] 🎉 翻译大功告成！双语文档已保存至:\n    -> {output_docx}")
    print("=" * 40)

if __name__ == '__main__':
    intput_file = "test_files/sample.pdf"
    output_file = "output/translated_result.docx"

    create_bilingual_word(intput_file, output_file)
